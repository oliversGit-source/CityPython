from docx import Document

filename = '/Users/oliverbutterworth-bakhshi/Desktop/Jogging.docx'
doc = Document(filename)
def getText(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        '\n'.join(fullText)
    print(f'Here is the text extracted from {filename}:')
    for x in fullText:
        print(x)

def wordcount(word):
    count = 0
    fullText = []
    foundlist = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        '\n'.join(fullText)
    for x in fullText:
        foundlist.append(x.find(word))
        count += 2
    print(f'The number of times {word} appeared in the search was {count}')

wordcount('heart')



